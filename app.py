from pydoc import text
from flask import Flask, request, render_template, redirect, send_from_directory, url_for, flash, jsonify
from flask import Flask, request, render_template, redirect, url_for, flash, session
from werkzeug.security import generate_password_hash, check_password_hash
import fitz  # PyMuPDF
import docx
import os
from dotenv import load_dotenv
import google.generativeai as genai
from db_config import get_db_connection
from flask_cors import CORS
import re
from elevenlabs import generate, save, set_api_key
import uuid
from docx import Document



# Load environment variables
load_dotenv()

# Gemini API setup
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Initialize Flask
app = Flask(__name__)
app.secret_key = 'your_secret_key'
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name")
        username = request.form.get("username")
        email = request.form.get("email")
        phone = request.form.get("phone")
        password = generate_password_hash(request.form.get("password"))

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO users (name, username, email, phone, password)
                VALUES (%s, %s, %s, %s, %s)
            """, (name, username, email, phone, password))
            conn.commit()
            cursor.close()
            conn.close()
            flash("✅ Registered successfully. Please log in.")
            return redirect("/login")
        except Exception as e:
            flash(f"❌ Registration failed: {str(e)}")

    return render_template("register.html")
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()
        cursor.close()
        conn.close()

        if user and check_password_hash(user["password"], password):
            session['user'] = user["username"]
            return redirect("/index")
        else:
            flash("❌ Invalid email or password.")

    return render_template("login.html")


def extract_text(file):
    try:
        if hasattr(file, 'read') and hasattr(file, 'filename') and file.filename:
            ext = os.path.splitext(file.filename)[1].lower()
            print(f"🧾 Reading stream file: {file.filename}")
            file.stream.seek(0)
            if ext == ".pdf":
                pdf = fitz.open(stream=file.read(), filetype="pdf")
                return "".join([page.get_text() for page in pdf])
            elif ext == ".docx":
                doc = docx.Document(file)
                return "\n".join([p.text for p in doc.paragraphs])
        elif isinstance(file, str) and os.path.exists(file):
            ext = os.path.splitext(file)[1].lower()
            print(f"📂 Reading saved path: {file}")
            if ext == ".pdf":
                with open(file, "rb") as f:
                    pdf = fitz.open(stream=f.read(), filetype="pdf")
                    return "".join([page.get_text() for page in pdf])
            elif ext == ".docx":
                doc = docx.Document(file)
                return "\n".join([p.text for p in doc.paragraphs])
        else:
            print("⚠️ File path or object invalid.")
    except Exception as e:
        print("❌ extract_text error:", e)
    return ""





# Analyze document with Gemini
def analyze_with_gemini(text, domain, policy):
    prompt = f"""
You are a Policy Compliance Analyst AI with in-depth knowledge of global policies like GDPR, ISO 27001, HIPAA, ACA, CPSIA, and others.

Analyze the internal policy below and compare it **only** with the selected real-world policy: "{policy}" under "{domain}" industry.

Focus on highlighting mismatches or gaps between the internal document and the selected real-world policy.

== Document Start ==
{text[:3000]}
== Document End ==

Respond ONLY in this structure:

Summary:
<Short paragraph>

Issues:
- <up to 4 compliance issues vs selected policy>

Recommendations:
- <up to 4 improvement suggestions>

Role-wise Guidelines:
- HR:
  * Policy Article: <Article number or Section Number>
    Current Policy: <one sentence>
    {policy} Requirement: <one sentence>
    Amendment Suggestion: <one sentence>
- IT:
  * Policy Article: <...>
- Finance:
  * Policy Article: <...>

Penalties:
- <Article reference and consequence>

Risk Level: <Low / Medium / High>
Effective Date: <DD-MM-YYYY or Month-Year>
Affiliated Article: <Most relevant article/section number>
Policy URL: <Link to official policy>
Owner: <Owner Name or "Unavailable">

Output must be plain text in this exact structure.
"""

    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash-lite-preview-06-17")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"ERROR: {str(e)}"

@app.route("/")
def redirect_to_home():
    return redirect("/home")

@app.route("/home", methods=["GET", "POST"])
def home():
    if request.method == "POST":
        name = request.form["name"]
        email = request.form["email"]
        phone = request.form.get("phone", "")
        place = request.form.get("place", "")
        message = request.form.get("message", "")

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO contact_messages (name, email, phone, place, message)
                VALUES (%s, %s, %s, %s, %s)
            """, (name, email, phone, place, message))
            conn.commit()
            cursor.close()
            conn.close()
            flash("✅ Your message has been submitted successfully!")
        except Exception as e:
            flash(f"❌ Something went wrong: {str(e)}")

        return redirect(url_for('home'))

    return render_template("home.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    file = request.files["file"]
    domain = request.form.get("industry", "General")
    policy = request.form.get("policy", "General")
    text = extract_text(file)
    result = analyze_with_gemini(text, domain, policy)

    if result.startswith("ERROR"):
        return render_template("analysis.html", analysis_summary=result, issues=[], recommendations=[], role_guidelines={}, penalties=[], owner="Not Available", domain=domain, policy=policy, risk_level="Not Available")


    try:
        summary = result.split("Issues:")[0].replace("Summary:", "").strip()
        issues = result.split("Issues:")[1].split("Recommendations:")[0].strip().split("- ")
        recs = result.split("Recommendations:")[1].split("Role-wise Guidelines:")[0].strip().split("- ")
        role_section = result.split("Role-wise Guidelines:")[1].split("Penalties:")[0]
        penalties_section = result.split("Penalties:")[1].split("Risk Level:")[0].strip()
        risk_level = result.split("Risk Level:")[1].split("Effective Date:")[0].strip()
        effective_date = result.split("Effective Date:")[1].split("Affiliated Article:")[0].strip()
        affiliated_article = result.split("Affiliated Article:")[1].split("Policy URL:")[0].strip()
        policy_url = result.split("Policy URL:")[1].split("Owner:")[0].strip()
        owner = result.split("Owner:")[-1].strip()
        unique_filename = f"{uuid.uuid4().hex}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(filepath)

        penalties = [p.strip("- ").strip() for p in penalties_section.strip().split("\n") if p.strip()]

        role_guidelines = {}
        for role in ["HR", "IT", "Finance"]:
            structured = []
            try:
                role_block = role_section.split(f"- {role}:")[1].split("- ")[0].strip()
                blocks = role_block.split("* Policy Article:")

                for block in blocks[1:]:
                    try:
                        article_part, rest = block.strip().split("Current Policy:")
                        article = re.search(r"(Article|Section)?\s*\d+[\w\.\(\)-]*", article_part.strip(), re.IGNORECASE)
                        article = article.group(0) if article else "N/A"

                        if "Requirement:" not in rest:
                            raise ValueError("Missing Requirement")
                        current, rest2 = rest.split("Requirement:")

                        if "Amendment Suggestion:" not in rest2:
                            raise ValueError("Missing Amendment")
                        requirement, amendment = rest2.split("Amendment Suggestion:")

                        structured.append({
                            "article": article[:100],
                            "current": current.strip()[:250] or "Not Available",
                            "requirement": requirement.strip()[:250] or "Not Available",
                            "amendment": amendment.strip()[:250] or "Not Available"
                        })
                    except Exception:
                        structured.append({
                            "article": "N/A",
                            "current": "Not Available",
                            "requirement": "Not Available",
                            "amendment": "Not Available"
                        })
            except Exception:
                structured = [{
                    "article": "N/A",
                    "current": "Not Available",
                    "requirement": "Not Available",
                    "amendment": "Not Available"
                }]
            role_guidelines[role] = structured

        return render_template("analysis.html",
                               analysis_summary=summary,
                               issues=[i.strip() for i in issues if i.strip()][:4],
                               recommendations=[r.strip() for r in recs if r.strip()][:4],
                               role_guidelines=role_guidelines,
                               penalties=penalties,
                               owner=owner,
                               domain=domain,
                               policy=policy,
                               risk_level=risk_level,
                               effective_date=effective_date,
                               affiliated_article=affiliated_article,
                               policy_url=policy_url if policy_url else "#",
                               uploaded_filename=unique_filename)

    except Exception as e:
        return render_template("analysis.html",
                               analysis_summary="Could not parse response correctly.",
                               issues=["Parsing error."],
                               recommendations=["Parsing error."],
                               role_guidelines={},
                               penalties=[],
                               owner="Unavailable",
                               domain=domain,
                               policy=policy,
                               risk_level="Unavailable",
                               effective_date="Unavailable",
                               affiliated_article="Unavailable",
                               policy_url="#",
                               uploaded_filename=file.filename)

@app.route("/generate_draft", methods=["POST"])
def generate_draft():
    data = request.get_json()
    file_name = data.get("file_name", "")
    policy = data.get("policy", "")

    try:
        filepath = os.path.join(UPLOAD_FOLDER, file_name)
        if not os.path.exists(filepath):
            raise FileNotFoundError("File not found")
        text = extract_text(filepath)
    except Exception as e:
        return jsonify({"draft": f"❌ Failed to load original document. {str(e)}"})

    prompt = f"""
Based on prior analysis and best practices for {policy} compliance, rewrite the internal policy below.

✅ Output ONLY the rewritten and amended version — do not explain what you will do.
✅ Format the result as a complete company internal policy document.
✅ Use clear headings like Purpose, Scope, Policy Statement, Candidate Rights, Governance, etc.
✅ Ensure it is formally worded and suitable for direct adoption by a company.

=== ORIGINAL POLICY START ===
{text[:3000]}
=== ORIGINAL POLICY END ===
"""

    
    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash-lite-preview-06-17")
        response = model.generate_content(prompt)
        return jsonify({"draft": response.text})
    except Exception as e:
        return jsonify({"draft": f"❌ Error: {str(e)}"})




@app.route("/index")
def index():
    if 'user' not in session:
        return redirect("/login")
    return render_template("index.html")
@app.route("/logout")
def logout():
    session.pop('user', None)
    return redirect("/home")

@app.route("/getstarted")
def getstarted():
    return redirect("/login")

# Optional: Smart Assistant voice synthesis (not supported natively in Flask, use front-end JS + ElevenLabs or Web Speech API)
@app.route("/chatbot", methods=["POST"])
def chatbot():
    data = request.get_json()
    user_message = data.get("message", "")

    if not user_message:
        return jsonify({"reply": "❌ Please enter a message."})

    prompt = f"""
You are a helpful and concise AI assistant specialized in global policy compliance (GDPR, ISO 27001, HIPAA, etc.).

Respond briefly and to the point. Avoid introductions, apologies, or repeating the question.

Limit your response to **under 40 words**.

User: "{user_message}"
Assistant:
"""


 
    try:
        model = genai.GenerativeModel("models/gemini-2.5-flash-lite-preview-06-17")
        response = model.generate_content(prompt)
        return jsonify({"reply": response.text})
    except Exception as e:
        return jsonify({"reply": f"❌ Error: {str(e)}"})
@app.route("/ask_queries")
def ask_queries():
    return render_template("ask_queries.html")  # Make sure you have this HTML file in templates/

@app.route("/speak", methods=["POST"])
def speak():
    from elevenlabs import generate, save, set_api_key
    import os
    data = request.get_json()
    text = data.get("text", "")[:300]
    try:
        set_api_key(os.getenv("ELEVENLABS_API_KEY"))

        # Create a unique filename
        filename = f"{uuid.uuid4().hex}.mp3"
        filepath = os.path.join("static", "audio", filename)

        # Generate voice audio
        audio = generate(
            text=text,
            voice="XrExE9yKIg1WjnnlVkGX",  # You can change this to any voice ID
            model="eleven_multilingual_v2"
        )

        # Save the audio file
        save(audio, filepath)

        return jsonify({ "audio_url": f"/static/audio/{filename}" })
    except Exception as e:
        return jsonify({ "error": str(e) }), 500

        # Save in static/audio so it's accessible
        audio_dir = os.path.join("static", "audio")
        os.makedirs(audio_dir, exist_ok=True)
        filename = f"audio_{uuid.uuid4().hex}.mp3"
        filepath = os.path.join(audio_dir, filename)

        save(audio, filepath)

        return jsonify({"audio_url": f"/static/audio/{filename}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route("/audio/<filename>")
def serve_audio(filename):
    return send_from_directory("/tmp", filename)

@app.route("/feedback")
def feedback():
    return render_template("feedback.html")

@app.route("/policy-guidelines")
def policy_guidelines():
    return render_template("policy_guidelines.html")
@app.route("/download_draft", methods=["POST"])
def download_draft():
    data = request.get_json()
    draft_text = data.get("draft_text", "")
    if not draft_text:
        return jsonify({"error": "❌ No draft content to download."}), 400

    # Create Word document
    document = Document()
    for line in draft_text.split("\n"):
        document.add_paragraph(line)

    # Save to temp path
    filename = f"{uuid.uuid4().hex}_draft.docx"
    path = os.path.join("static", "downloads", filename)
    os.makedirs("static/downloads", exist_ok=True)
    document.save(path)

    return jsonify({ "download_url": f"/static/downloads/{filename}" })



@app.route("/submit_query", methods=["POST"])
def submit_query():
    name = request.form["name"]
    email = request.form["email"]
    subject = request.form["subject"]
    message = request.form["message"]

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO ask_queries (name, email, subject, message)
            VALUES (%s, %s, %s, %s)
        """, (name, email, subject, message))
        conn.commit()
        cursor.close()
        conn.close()
        flash("✅ Your query has been submitted successfully!")
    except Exception as e:
        flash(f"❌ Error submitting query: {str(e)}")

    return redirect(url_for("ask_queries"))


@app.route("/submit_feedback", methods=["POST"])
def submit_feedback():
    name = request.form["name"]
    email = request.form["email"]
    rating = request.form["rating"]
    comments = request.form["comments"]

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO feedback (name, email, rating, comments)
            VALUES (%s, %s, %s, %s)
        """, (name, email, rating, comments))
        conn.commit()
        cursor.close()
        conn.close()
        flash("⭐ Thank you for your feedback!")
    except Exception as e:
        flash(f"❌ Error: {str(e)}")

    return redirect(url_for("feedback"))

if __name__ == "__main__":
    app.run(debug=True)
